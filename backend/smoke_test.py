"""Smoke test for the Re:mind MVP V0 FastAPI backend.

Run from the backend directory:
    uv run python smoke_test.py
"""
from __future__ import annotations

import contextlib
import json
import os
import re
import tempfile
import time
import zipfile
from io import BytesIO
from pathlib import Path
from urllib.parse import unquote

from fastapi.testclient import TestClient

from app.core.config import settings, validate_runtime_security
from app.graph import nodes as graph_nodes
from app.main import app
from app.schemas.note import (
    ConfirmGeneratedNoteRequest,
    GenerateNoteResponse,
    RetrievedCaseContextItem,
    RetrievedEvidenceItem,
    RetrievedPrivacyRule,
    RetrievedTemplateContext,
    SessionInput,
)
from app.services.deidentification import deidentify_text
from app.services.embeddings import (
    EmbeddingError,
    clear_embedding_cache,
    content_hash,
    embed_query,
    embedding_cache_stats,
    get_embedding_provider,
)
from app.services.retrieval import RetrievalChunk
from app.services import retrieval as retrieval_module
from app.services.retrieval import retrieve_case_memory_chunks
from app.services import supabase_storage as supabase_storage_module
from app.services.supabase_storage import (
    ConfirmedNoteContext,
    NoteConfirmationError,
    _attach_embeddings,
    _build_session_row,
    _case_memory_rows_from_confirmed_note,
    confirm_generated_note,
)


def _extract_docx_text(content: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells if cell.text)
    return "\n".join(parts)


def _download_filename(content_disposition: str) -> str:
    match = re.search(r"filename\*=UTF-8''([^;]+)", content_disposition)
    if match:
        return unquote(match.group(1))
    fallback = re.search(r'filename="?([^";]+)"?', content_disposition)
    return fallback.group(1) if fallback else ""


def _assert_pdf_response(content: bytes, content_type: str, expected_texts: list[str]) -> None:
    from pypdf import PdfReader

    assert content.startswith(b"%PDF")
    assert content_type.startswith("application/pdf")
    reader = PdfReader(BytesIO(content))
    assert len(reader.pages) >= 1
    extracted_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    if extracted_text.strip() and any(expected in extracted_text for expected in expected_texts):
        return
    if extracted_text.strip():
        print("PDF text extraction did not preserve expected Korean text; page-count validation passed.")


def _make_docx_bytes() -> bytes:
    from docx import Document

    document = Document()
    document.add_paragraph("상담 메모 첫 문단")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "영역"
    table.cell(0, 1).text = "내용"
    table.cell(1, 0).text = "정서"
    table.cell(1, 1).text = "불안 80"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_text_pdf_bytes(text: str = "Hello counseling note PDF text") -> bytes:
    safe_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({safe_text}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    content = b"%PDF-1.4\n"
    offsets = []
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(content))
        content += f"{index} 0 obj\n".encode("ascii") + obj + b"\nendobj\n"
    xref_offset = len(content)
    content += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii")
    for offset in offsets:
        content += f"{offset:010d} 00000 n \n".encode("ascii")
    content += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    ).encode("ascii")
    return content


def _make_blank_pdf_bytes() -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def _make_encrypted_pdf_bytes() -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.encrypt("secret")
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def _make_zip_bytes(entries: dict[str, bytes]) -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, content in entries.items():
            archive.writestr(name, content)
    return buffer.getvalue()


def _make_wav_bytes(payload_size: int = 16) -> bytes:
    data = b"\x00" * payload_size
    return (
        b"RIFF"
        + (36 + len(data)).to_bytes(4, "little")
        + b"WAVEfmt "
        + (16).to_bytes(4, "little")
        + (1).to_bytes(2, "little")
        + (1).to_bytes(2, "little")
        + (16000).to_bytes(4, "little")
        + (32000).to_bytes(4, "little")
        + (2).to_bytes(2, "little")
        + (16).to_bytes(2, "little")
        + b"data"
        + len(data).to_bytes(4, "little")
        + data
    )


def _make_mp3_bytes() -> bytes:
    return b"ID3\x04\x00\x00\x00\x00\x00\x10" + (b"\x00" * 32)


def _make_m4a_bytes() -> bytes:
    return b"\x00\x00\x00\x18ftypM4A \x00\x00\x00\x00M4A mp42"


def _upload_file(client: TestClient, name: str, body: bytes, content_type: str):
    return client.post(
        "/api/materials/documents/extract",
        files={"file": (name, BytesIO(body), content_type)},
    )


def _assert_signature_checks_are_streaming() -> None:
    from app.services.upload_validation import signature_matches

    samples = [
        (".wav", "audio", _make_wav_bytes(1024)),
        (".mp3", "audio", _make_mp3_bytes()),
        (".m4a", "audio", _make_m4a_bytes()),
        (".pdf", "document", _make_text_pdf_bytes()),
    ]

    original_read_bytes = Path.read_bytes
    original_open = Path.open
    read_sizes: list[int] = []

    def fail_read_bytes(self: Path) -> bytes:
        raise AssertionError("Path.read_bytes() must not be used for upload signature checks")

    class GuardedFile:
        def __init__(self, handle):
            self.handle = handle

        def __enter__(self):
            self.handle.__enter__()
            return self

        def __exit__(self, exc_type, exc, traceback):
            return self.handle.__exit__(exc_type, exc, traceback)

        def read(self, size: int = -1) -> bytes:
            read_sizes.append(size)
            assert 0 <= size <= 64, f"signature check read too much: {size}"
            return self.handle.read(size)

        def __getattr__(self, name: str):
            return getattr(self.handle, name)

    with tempfile.TemporaryDirectory() as temp_dir:
        watched_paths: set[Path] = set()
        for suffix, kind, content in samples:
            sample_path = Path(temp_dir) / f"sample{suffix}"
            sample_path.write_bytes(content)
            watched_paths.add(sample_path)

        def guarded_open(self: Path, *args, **kwargs):
            handle = original_open(self, *args, **kwargs)
            if self in watched_paths:
                return GuardedFile(handle)
            return handle

        try:
            Path.read_bytes = fail_read_bytes
            Path.open = guarded_open
            for suffix, kind, _content in samples:
                assert signature_matches(Path(temp_dir) / f"sample{suffix}", suffix, kind)
        finally:
            Path.read_bytes = original_read_bytes
            Path.open = original_open

    assert read_sizes and all(size == 64 for size in read_sizes)


def _run_material_upload_smoke_tests(client: TestClient) -> None:
    _assert_signature_checks_are_streaming()

    txt_response = _upload_file(
        client,
        "memo.txt",
        "\ufeff상담 메모\n둘째 줄".encode("utf-8"),
        "text/plain",
    )
    assert txt_response.status_code == 200, txt_response.text
    txt_data = txt_response.json()
    assert txt_data["status"] == "completed"
    assert "상담 메모" in txt_data["extracted_text"]
    assert "둘째 줄" in txt_data["extracted_text"]
    assert txt_data["character_count"] >= 8
    assert txt_response.headers["cache-control"] == "no-store"

    docx_response = _upload_file(
        client,
        "case-note.docx",
        _make_docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert docx_response.status_code == 200, docx_response.text
    docx_text = docx_response.json()["extracted_text"]
    assert "상담 메모 첫 문단" in docx_text
    assert "영역\t내용" in docx_text
    assert "정서\t불안 80" in docx_text

    pdf_response = _upload_file(client, "note.pdf", _make_text_pdf_bytes(), "application/pdf")
    assert pdf_response.status_code == 200, pdf_response.text
    pdf_data = pdf_response.json()
    assert pdf_data["page_count"] == 1
    assert "Hello counseling note PDF text" in pdf_data["extracted_text"]

    image_pdf_response = _upload_file(client, "scan.pdf", _make_blank_pdf_bytes(), "application/pdf")
    assert image_pdf_response.status_code == 200, image_pdf_response.text
    image_pdf_data = image_pdf_response.json()
    assert image_pdf_data["status"] == "warning"
    assert "OCR" in " ".join(image_pdf_data["warnings"])

    corrupt_docx_response = _upload_file(
        client,
        "broken.docx",
        b"PK\x03\x04not a readable zip archive",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert corrupt_docx_response.status_code == 422
    assert "DOCX 파일을 읽을 수 없습니다" in corrupt_docx_response.text

    corrupt_pdf_response = _upload_file(client, "broken.pdf", b"%PDF-1.4\nbroken", "application/pdf")
    assert corrupt_pdf_response.status_code == 422
    assert "PDF 파일을 읽을 수 없습니다" in corrupt_pdf_response.text

    encrypted_pdf_response = _upload_file(client, "encrypted.pdf", _make_encrypted_pdf_bytes(), "application/pdf")
    assert encrypted_pdf_response.status_code == 422
    assert "암호화된 PDF" in encrypted_pdf_response.text

    original_docx_member_limit = os.environ.get("DOCX_MAX_ARCHIVE_MEMBERS")
    os.environ["DOCX_MAX_ARCHIVE_MEMBERS"] = "2"
    try:
        too_many_members = _upload_file(
            client,
            "too-many.docx",
            _make_zip_bytes(
                {
                    "[Content_Types].xml": b"<Types/>",
                    "word/document.xml": b"<document/>",
                    "word/extra.xml": b"<extra/>",
                }
            ),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert too_many_members.status_code == 413
    finally:
        if original_docx_member_limit is None:
            os.environ.pop("DOCX_MAX_ARCHIVE_MEMBERS", None)
        else:
            os.environ["DOCX_MAX_ARCHIVE_MEMBERS"] = original_docx_member_limit

    original_docx_size_limit = os.environ.get("DOCX_MAX_UNCOMPRESSED_BYTES")
    os.environ["DOCX_MAX_UNCOMPRESSED_BYTES"] = "30"
    try:
        too_large_docx = _upload_file(
            client,
            "too-large.docx",
            _make_zip_bytes(
                {
                    "[Content_Types].xml": b"<Types/>",
                    "word/document.xml": b"<document/>",
                    "word/large.bin": b"x" * 64,
                }
            ),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert too_large_docx.status_code == 413
    finally:
        if original_docx_size_limit is None:
            os.environ.pop("DOCX_MAX_UNCOMPRESSED_BYTES", None)
        else:
            os.environ["DOCX_MAX_UNCOMPRESSED_BYTES"] = original_docx_size_limit

    original_docx_ratio_limit = os.environ.get("DOCX_MAX_COMPRESSION_RATIO")
    os.environ["DOCX_MAX_COMPRESSION_RATIO"] = "2"
    try:
        high_ratio_docx = _upload_file(
            client,
            "high-ratio.docx",
            _make_zip_bytes(
                {
                    "[Content_Types].xml": b"<Types/>",
                    "word/document.xml": b"<document/>",
                    "word/repeated.txt": b"a" * 4096,
                }
            ),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert high_ratio_docx.status_code == 413
    finally:
        if original_docx_ratio_limit is None:
            os.environ.pop("DOCX_MAX_COMPRESSION_RATIO", None)
        else:
            os.environ["DOCX_MAX_COMPRESSION_RATIO"] = original_docx_ratio_limit

    unsupported_response = _upload_file(client, "note.rtf", b"{\\rtf1}", "application/rtf")
    assert unsupported_response.status_code == 415

    mismatch_response = _upload_file(client, "note.pdf", _make_text_pdf_bytes(), "text/plain")
    assert mismatch_response.status_code == 415

    empty_response = _upload_file(client, "empty.txt", b"", "text/plain")
    assert empty_response.status_code == 400

    original_doc_limit = os.environ.get("DOCUMENT_UPLOAD_MAX_BYTES")
    os.environ["DOCUMENT_UPLOAD_MAX_BYTES"] = "8"
    try:
        too_large_response = _upload_file(client, "large.txt", b"0123456789", "text/plain")
        assert too_large_response.status_code == 413
    finally:
        if original_doc_limit is None:
            os.environ.pop("DOCUMENT_UPLOAD_MAX_BYTES", None)
        else:
            os.environ["DOCUMENT_UPLOAD_MAX_BYTES"] = original_doc_limit

    with tempfile.TemporaryDirectory() as upload_temp_dir:
        previous_tmp_dir = os.environ.get("UPLOAD_TMP_DIR")
        os.environ["UPLOAD_TMP_DIR"] = upload_temp_dir
        try:
            cleanup_response = _upload_file(client, "cleanup.txt", b"temporary cleanup text", "text/plain")
            assert cleanup_response.status_code == 200, cleanup_response.text
            assert not list(Path(upload_temp_dir).iterdir())
        finally:
            if previous_tmp_dir is None:
                os.environ.pop("UPLOAD_TMP_DIR", None)
            else:
                os.environ["UPLOAD_TMP_DIR"] = previous_tmp_dir

    with tempfile.TemporaryDirectory() as upload_temp_dir:
        previous_tmp_dir = os.environ.get("UPLOAD_TMP_DIR")
        os.environ["UPLOAD_TMP_DIR"] = upload_temp_dir
        try:
            cleanup_after_parser_failure = _upload_file(
                client,
                "cleanup-broken.pdf",
                b"%PDF-1.4\nbroken",
                "application/pdf",
            )
            assert cleanup_after_parser_failure.status_code == 422
            assert not list(Path(upload_temp_dir).iterdir())
        finally:
            if previous_tmp_dir is None:
                os.environ.pop("UPLOAD_TMP_DIR", None)
            else:
                os.environ["UPLOAD_TMP_DIR"] = previous_tmp_dir

    stdout_buffer = BytesIO()
    stderr_buffer = BytesIO()
    secret_text = "RAW_SECRET_CONTENT_SHOULD_NOT_BE_LOGGED"
    with contextlib.redirect_stdout(_BytesTextWriter(stdout_buffer)), contextlib.redirect_stderr(
        _BytesTextWriter(stderr_buffer)
    ):
        raw_log_response = _upload_file(client, "secret.txt", secret_text.encode("utf-8"), "text/plain")
    assert raw_log_response.status_code == 200, raw_log_response.text
    assert secret_text.encode("utf-8") not in stdout_buffer.getvalue()
    assert secret_text.encode("utf-8") not in stderr_buffer.getvalue()

    capabilities = client.get("/api/audio/capabilities")
    assert capabilities.status_code == 200, capabilities.text
    audio_capabilities = capabilities.json()
    assert audio_capabilities["upload"]["available"] is True
    assert audio_capabilities["transcription"]["available"] is False
    assert audio_capabilities["speaker_diarization"]["available"] is False
    assert audio_capabilities["runtime_mode"] == "disabled"

    original_audio_limit = os.environ.get("AUDIO_UPLOAD_MAX_BYTES")
    os.environ["AUDIO_UPLOAD_MAX_BYTES"] = "20"
    try:
        audio_too_large = client.post(
            "/api/audio/transcribe",
            files={"file": ("session.wav", BytesIO(_make_wav_bytes(128)), "audio/wav")},
        )
        assert audio_too_large.status_code == 413
    finally:
        if original_audio_limit is None:
            os.environ.pop("AUDIO_UPLOAD_MAX_BYTES", None)
        else:
            os.environ["AUDIO_UPLOAD_MAX_BYTES"] = original_audio_limit

    audio_unavailable = client.post(
        "/api/audio/transcribe",
        files={"file": ("session.wav", BytesIO(_make_wav_bytes()), "audio/wav")},
        data={"language": "ko", "task": "transcribe"},
    )
    assert audio_unavailable.status_code == 503

    invalid_task = client.post(
        "/api/audio/transcribe",
        files={"file": ("session.wav", BytesIO(_make_wav_bytes()), "audio/wav")},
        data={"language": "ko", "task": "summarize"},
    )
    assert invalid_task.status_code == 422

    invalid_language = client.post(
        "/api/audio/transcribe",
        files={"file": ("session.wav", BytesIO(_make_wav_bytes()), "audio/wav")},
        data={"language": "../secret", "task": "transcribe"},
    )
    assert invalid_language.status_code == 422

    invalid_speaker_count = client.post(
        "/api/audio/transcribe",
        files={"file": ("session.wav", BytesIO(_make_wav_bytes()), "audio/wav")},
        data={"language": "ko", "task": "transcribe", "expected_speakers": "5"},
    )
    assert invalid_speaker_count.status_code == 422

    _run_audio_transcription_runtime_smoke_tests(client)


def _run_audio_transcription_runtime_smoke_tests(client: TestClient) -> None:
    from app.schemas.audio import AudioSegment
    from app.services import audio_transcription as audio_service
    from app.services.upload_validation import ValidatedUpload

    env_names = [
        "ENABLE_AUDIO_TRANSCRIPTION",
        "AUDIO_TRANSCRIPTION_STUB",
        "AUDIO_TRANSCRIPTION_ENGINE",
        "WHISPERX_MODEL",
        "WHISPERX_LANGUAGE",
        "WHISPERX_DEVICE",
        "WHISPERX_COMPUTE_TYPE",
        "WHISPERX_BATCH_SIZE",
        "WHISPERX_ALIGN_MODEL",
        "ENABLE_AUDIO_DIARIZATION",
        "WHISPERX_DIARIZATION_MODEL",
        "HF_TOKEN",
        "AUDIO_MAX_DURATION_SECONDS",
        "AUDIO_MAX_CONCURRENT_JOBS",
    ]
    previous_env = {name: os.environ.get(name) for name in env_names}
    os.environ.update(
        {
            "ENABLE_AUDIO_TRANSCRIPTION": "1",
            "AUDIO_TRANSCRIPTION_STUB": "0",
            "AUDIO_TRANSCRIPTION_ENGINE": "whisperx",
            "WHISPERX_MODEL": "large-v3",
            "WHISPERX_LANGUAGE": "ko",
            "WHISPERX_DEVICE": "auto",
            "WHISPERX_COMPUTE_TYPE": "float16",
            "WHISPERX_BATCH_SIZE": "4",
            "WHISPERX_ALIGN_MODEL": "kresnik/wav2vec2-large-xlsr-korean",
            "ENABLE_AUDIO_DIARIZATION": "1",
            "WHISPERX_DIARIZATION_MODEL": "pyannote/speaker-diarization-community-1",
            "HF_TOKEN": "hf-test-token",
            "AUDIO_MAX_DURATION_SECONDS": "7200",
            "AUDIO_MAX_CONCURRENT_JOBS": "1",
        }
    )

    class FakeCuda:
        def __init__(self, available: bool) -> None:
            self.available = available

        def is_available(self) -> bool:
            return self.available

    class FakeTorch:
        def __init__(self, cuda_available: bool) -> None:
            self.cuda = FakeCuda(cuda_available)

    class FakeAsrModel:
        def __init__(self, owner: "FakeWhisperX") -> None:
            self.owner = owner

        def transcribe(self, audio, **kwargs):
            self.owner.calls.append(("asr_transcribe", kwargs))
            if self.owner.fail_asr:
                raise RuntimeError("unsafe asr failure with server details")
            return {
                "language": self.owner.detected_language,
                "language_probability": 0.98,
                "segments": [
                    {
                        "start": 0.0,
                        "end": 2.2,
                        "text": "첫 번째 발화 상담자 반영",
                        "words": [
                            {"start": 0.0, "end": 0.4, "word": "첫", "score": 0.91},
                            {"start": 0.4, "end": 0.9, "word": "번째 발화", "score": 0.92},
                            {"start": 1.2, "end": 1.7, "word": "상담자", "score": 0.93},
                            {"start": 1.7, "end": 2.2, "word": "반영", "score": 0.94},
                        ],
                    }
                ],
            }

    class FakeWhisperX:
        def __init__(self) -> None:
            self.calls: list[tuple[str, object]] = []
            self.audio = [0.05] * 30
            self.detected_language = "ko"
            self.fail_asr = False
            self.fail_alignment = False
            self.fail_diarization = False

        def load_audio(self, path: str):
            self.calls.append(("load_audio", path))
            return self.audio

        def load_model(self, model_name: str, device: str, **kwargs):
            self.calls.append(("load_model", {"model_name": model_name, "device": device, **kwargs}))
            return FakeAsrModel(self)

        def load_align_model(self, **kwargs):
            self.calls.append(("load_align_model", kwargs))
            return object(), {"language": kwargs["language_code"]}

        def align(self, segments, model, metadata, audio, device, **kwargs):
            self.calls.append(("align", {"segments": segments, "device": device, **kwargs}))
            if self.fail_alignment:
                raise RuntimeError("unsafe alignment failure")
            return {"language": self.detected_language, "segments": segments}

        def assign_word_speakers(self, diarization_result, aligned_result):
            self.calls.append(("assign_word_speakers", diarization_result))
            assigned_segments = []
            for segment in aligned_result["segments"]:
                words = []
                for word in segment.get("words", []):
                    speaker = "SPEAKER_00" if word["start"] < 1.1 else "SPEAKER_01"
                    words.append({**word, "speaker": speaker})
                assigned_segments.append(
                    {
                        **segment,
                        "speaker": words[0]["speaker"] if words else "SPEAKER_00",
                        "words": words,
                    }
                )
            return {"language": self.detected_language, "segments": assigned_segments}

    class FakeDiarizationPipeline:
        def __init__(self, owner: FakeWhisperX) -> None:
            self.owner = owner

        def __call__(self, audio, **kwargs):
            self.owner.calls.append(("diarization_call", kwargs))
            if self.owner.fail_diarization:
                raise RuntimeError("unsafe diarization failure")
            return {"speaker_turns": "mock"}

    fake_whisperx = FakeWhisperX()

    def fake_diarization_factory(**kwargs):
        fake_whisperx.calls.append(("diarization_factory", kwargs))
        return FakeDiarizationPipeline(fake_whisperx)

    runtime = audio_service.WhisperXRuntime(
        whisperx=fake_whisperx,
        diarization_pipeline_factory=fake_diarization_factory,
        torch=FakeTorch(cuda_available=True),
        sample_rate=10,
    )

    def make_validated_audio() -> tuple[ValidatedUpload, Path]:
        temp_file = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
        temp_file.write(_make_wav_bytes())
        temp_file.close()
        path = Path(temp_file.name)
        return (
            ValidatedUpload(
                filename="session.wav",
                media_type="audio/wav",
                suffix=".wav",
                size_bytes=path.stat().st_size,
                temp_path=path,
            ),
            path,
        )

    try:
        audio_service.set_whisperx_runtime_for_testing(runtime)
        first_service = audio_service.get_transcription_service()
        second_service = audio_service.get_transcription_service()
        assert first_service is second_service

        validated, audio_path = make_validated_audio()
        try:
            first_result = first_service.transcribe(
                validated,
                language="ko",
                task="transcribe",
                expected_speakers=2,
            )
            second_result = second_service.transcribe(
                validated,
                language="ko",
                task="transcribe",
                expected_speakers=2,
            )
        finally:
            audio_path.unlink(missing_ok=True)

        assert first_result.runtime_mode == "real"
        assert first_result.transcription_engine == "whisperx"
        assert first_result.alignment_model == "kresnik/wav2vec2-large-xlsr-korean"
        assert first_result.diarization_model == "pyannote/speaker-diarization-community-1"
        assert first_result.alignment_status == "completed"
        assert first_result.diarization_status == "completed"
        assert len(first_result.segments) == 2
        assert [segment.speaker for segment in first_result.segments] == ["SPEAKER_00", "SPEAKER_01"]
        assert [word.speaker for segment in first_result.segments for word in segment.words] == [
            "SPEAKER_00",
            "SPEAKER_00",
            "SPEAKER_01",
            "SPEAKER_01",
        ]
        assert sum(name == "load_model" for name, _ in fake_whisperx.calls) == 1
        assert sum(name == "load_align_model" for name, _ in fake_whisperx.calls) == 1
        assert sum(name == "diarization_factory" for name, _ in fake_whisperx.calls) == 1

        load_model_call = next(data for name, data in fake_whisperx.calls if name == "load_model")
        assert load_model_call == {
            "model_name": "large-v3",
            "device": "cuda",
            "compute_type": "float16",
            "language": "ko",
            "task": "transcribe",
        }
        transcribe_call = next(data for name, data in fake_whisperx.calls if name == "asr_transcribe")
        assert transcribe_call["batch_size"] == 4
        assert transcribe_call["language"] == "ko"
        align_model_call = next(data for name, data in fake_whisperx.calls if name == "load_align_model")
        assert align_model_call == {
            "language_code": "ko",
            "device": "cuda",
            "model_name": "kresnik/wav2vec2-large-xlsr-korean",
        }
        align_call = next(data for name, data in fake_whisperx.calls if name == "align")
        assert align_call["return_char_alignments"] is False
        diarization_factory_call = next(
            data for name, data in fake_whisperx.calls if name == "diarization_factory"
        )
        assert diarization_factory_call == {
            "model_name": "pyannote/speaker-diarization-community-1",
            "token": "hf-test-token",
            "device": "cuda",
        }
        diarization_call = next(data for name, data in fake_whisperx.calls if name == "diarization_call")
        assert diarization_call["num_speakers"] == 2
        assert any(name == "assign_word_speakers" for name, _ in fake_whisperx.calls)

        gap_turns = audio_service._normalize_whisperx_turns(
            {
                "segments": [
                    {
                        "start": 0.0,
                        "end": 2.0,
                        "speaker": "SPEAKER_00",
                        "words": [
                            {"start": 0.0, "end": 0.2, "word": "첫 발화", "speaker": "SPEAKER_00"},
                            {
                                "start": 1.3,
                                "end": 1.8,
                                "word": "긴 간격 뒤 발화",
                                "speaker": "SPEAKER_00",
                            },
                        ],
                    }
                ]
            }
        )
        assert len(gap_turns) == 2

        acoustic_segments = [
            AudioSegment(id=1, start=0.0, end=1.0, text="천천히", speaker="SPEAKER_00"),
            AudioSegment(id=2, start=1.5, end=2.5, text="보통 속도", speaker="SPEAKER_00"),
            AudioSegment(id=3, start=2.6, end=3.0, text="아주 빠른 말 속도", speaker="SPEAKER_00"),
        ]
        acoustic_audio = ([0.01] * 10) + ([0.1] * 5) + ([0.1] * 10) + ([1.0] * 5)
        acoustic_result = audio_service._apply_deterministic_audio_features(
            acoustic_segments,
            acoustic_audio,
            10,
        )
        assert acoustic_result[1].pause_before_seconds == 0.5
        assert acoustic_result[0].speech_rate_level == "slow"
        assert acoustic_result[2].speech_rate_level == "fast"
        assert acoustic_result[0].volume_level == "low"
        assert acoustic_result[2].volume_level == "high"
        assert first_result.segments[0].speech_rate_level is None
        assert first_result.segments[0].volume_level is None

        fake_whisperx.fail_alignment = True
        os.environ["ENABLE_AUDIO_DIARIZATION"] = "0"
        audio_service.reset_transcription_service_cache_for_testing()
        validated, audio_path = make_validated_audio()
        try:
            alignment_fallback = audio_service.get_transcription_service().transcribe(validated)
        finally:
            audio_path.unlink(missing_ok=True)
        assert alignment_fallback.alignment_status == "fallback"
        assert alignment_fallback.transcript_text
        assert "한국어 단어 정렬에 실패" in " ".join(alignment_fallback.warnings)
        fake_whisperx.fail_alignment = False

        fake_whisperx.fail_diarization = True
        os.environ["ENABLE_AUDIO_DIARIZATION"] = "1"
        audio_service.set_whisperx_runtime_for_testing(runtime)
        validated, audio_path = make_validated_audio()
        try:
            diarization_fallback = audio_service.get_transcription_service().transcribe(
                validated,
                expected_speakers=2,
            )
        finally:
            audio_path.unlink(missing_ok=True)
        assert diarization_fallback.diarization_status == "fallback"
        assert {segment.speaker for segment in diarization_fallback.segments} == {"SPEAKER_00"}
        assert "화자 분리를 사용할 수 없어" in " ".join(diarization_fallback.warnings)
        assert "hf-test-token" not in diarization_fallback.model_dump_json()
        fake_whisperx.fail_diarization = False

        os.environ.pop("HF_TOKEN", None)
        audio_service.reset_transcription_service_cache_for_testing()
        validated, audio_path = make_validated_audio()
        try:
            missing_token_fallback = audio_service.get_transcription_service().transcribe(validated)
        finally:
            audio_path.unlink(missing_ok=True)
        assert missing_token_fallback.diarization_status == "fallback"
        assert missing_token_fallback.transcript_text
        assert {segment.speaker for segment in missing_token_fallback.segments} == {"SPEAKER_00"}
        os.environ["HF_TOKEN"] = "hf-test-token"

        fake_whisperx.detected_language = "en"
        audio_service.reset_transcription_service_cache_for_testing()
        validated, audio_path = make_validated_audio()
        try:
            language_mismatch = audio_service.get_transcription_service().transcribe(validated)
        finally:
            audio_path.unlink(missing_ok=True)
        assert "인식된 언어가 한국어 설정과 일치하지 않아" in " ".join(language_mismatch.warnings)
        fake_whisperx.detected_language = "ko"

        cpu_whisperx = FakeWhisperX()
        cpu_runtime = audio_service.WhisperXRuntime(
            whisperx=cpu_whisperx,
            diarization_pipeline_factory=lambda **kwargs: FakeDiarizationPipeline(cpu_whisperx),
            torch=FakeTorch(cuda_available=False),
            sample_rate=10,
        )
        os.environ["ENABLE_AUDIO_DIARIZATION"] = "0"
        audio_service.set_whisperx_runtime_for_testing(cpu_runtime)
        validated, audio_path = make_validated_audio()
        try:
            audio_service.get_transcription_service().transcribe(validated)
        finally:
            audio_path.unlink(missing_ok=True)
        cpu_load_call = next(data for name, data in cpu_whisperx.calls if name == "load_model")
        assert cpu_load_call["device"] == "cpu"
        assert cpu_load_call["compute_type"] == "int8"

        os.environ["AUDIO_MAX_DURATION_SECONDS"] = "1"
        audio_service.set_whisperx_runtime_for_testing(runtime)
        validated, audio_path = make_validated_audio()
        try:
            try:
                audio_service.get_transcription_service().transcribe(validated)
                raise AssertionError("duration limit must reject long audio")
            except audio_service.AudioDurationLimitError as error:
                assert str(error) == "음성 길이가 허용된 최대 처리 시간을 초과했습니다."
        finally:
            audio_path.unlink(missing_ok=True)
        os.environ["AUDIO_MAX_DURATION_SECONDS"] = "7200"

        audio_service.reset_transcription_service_cache_for_testing()
        busy_service = audio_service.get_transcription_service()
        semaphore = audio_service._get_audio_job_semaphore(1)
        assert semaphore.acquire(blocking=False)
        validated, audio_path = make_validated_audio()
        try:
            try:
                busy_service.transcribe(validated)
                raise AssertionError("concurrent job guard must reject excess work")
            except audio_service.AudioTranscriptionBusyError as error:
                assert str(error) == "다른 음성 축어록 작업이 진행 중입니다. 잠시 후 다시 시도해주세요."
        finally:
            semaphore.release()
            audio_path.unlink(missing_ok=True)

        fake_whisperx.fail_asr = True
        audio_service.set_whisperx_runtime_for_testing(runtime)
        with tempfile.TemporaryDirectory() as upload_temp_dir:
            previous_tmp_dir = os.environ.get("UPLOAD_TMP_DIR")
            os.environ["UPLOAD_TMP_DIR"] = upload_temp_dir
            try:
                route_failure = client.post(
                    "/api/audio/transcribe",
                    files={"file": ("session.wav", BytesIO(_make_wav_bytes()), "audio/wav")},
                    data={"language": "ko", "task": "transcribe", "expected_speakers": "2"},
                )
                assert route_failure.status_code == 500, route_failure.text
                assert route_failure.json()["detail"] == "음성 축어록 생성 중 오류가 발생했습니다."
                assert upload_temp_dir not in route_failure.text
                assert "unsafe asr failure" not in route_failure.text
                assert not list(Path(upload_temp_dir).iterdir())
            finally:
                if previous_tmp_dir is None:
                    os.environ.pop("UPLOAD_TMP_DIR", None)
                else:
                    os.environ["UPLOAD_TMP_DIR"] = previous_tmp_dir
        fake_whisperx.fail_asr = False

        calls_before_stub = len(fake_whisperx.calls)
        os.environ["AUDIO_TRANSCRIPTION_STUB"] = "1"
        os.environ["ENABLE_AUDIO_TRANSCRIPTION"] = "0"
        audio_service.reset_transcription_service_cache_for_testing()
        stub_response = client.post(
            "/api/audio/transcribe",
            files={"file": ("session.wav", BytesIO(_make_wav_bytes()), "audio/wav")},
            data={"language": "ko", "task": "transcribe", "expected_speakers": "3"},
        )
        assert stub_response.status_code == 200, stub_response.text
        stub_data = stub_response.json()
        assert stub_data["runtime_mode"] == "stub"
        assert stub_data["transcription_engine"] == "stub"
        assert "시연용 예시 축어록" in " ".join(stub_data["warnings"])
        assert {segment["speaker"] for segment in stub_data["segments"]} == {
            "speaker_1",
            "speaker_2",
            "speaker_3",
        }
        assert len(fake_whisperx.calls) == calls_before_stub
    finally:
        audio_service.set_whisperx_runtime_for_testing(None)
        for name, value in previous_env.items():
            if value is None:
                os.environ.pop(name, None)
            else:
                os.environ[name] = value

class _BytesTextWriter:
    def __init__(self, buffer: BytesIO) -> None:
        self.buffer = buffer

    def write(self, value: str) -> int:
        data = value.encode("utf-8", errors="replace")
        self.buffer.write(data)
        return len(value)

    def flush(self) -> None:
        return None


def main() -> None:
    require_pdf_export = os.getenv("REQUIRE_PDF_EXPORT") == "1"
    os.environ["ENABLE_AUDIO_TRANSCRIPTION"] = "0"
    os.environ["AUDIO_TRANSCRIPTION_STUB"] = "0"
    os.environ["ENABLE_AUDIO_DIARIZATION"] = "0"
    settings.use_stub = True
    settings.openai_api_key = None
    settings.enable_persistence = False
    settings.enable_rag = False
    settings.supabase_url = None
    settings.supabase_service_role_key = None
    settings.supabase_service_key = None
    settings.save_raw_input = False
    settings.enable_case_memory = False
    settings.runtime_environment = "test"
    settings.remind_preview_api_token = "test-preview-token"
    settings.remind_preview_actor = "test-preview-actor"
    settings.remind_allow_local_bypass = False
    settings.allow_legacy_preview_token = True

    original_runtime = settings.runtime_environment
    original_persistence = settings.enable_persistence
    original_case_memory = settings.enable_case_memory
    original_preview_token = settings.remind_preview_api_token
    original_real_auth = settings.enable_real_user_auth
    original_local_bypass = settings.remind_allow_local_bypass
    try:
        settings.runtime_environment = "production"
        settings.enable_persistence = True
        settings.enable_case_memory = False
        settings.remind_preview_api_token = None
        settings.enable_real_user_auth = False
        settings.remind_allow_local_bypass = False
        try:
            validate_runtime_security()
        except RuntimeError as error:
            assert "ENABLE_PERSISTENCE=1" in str(error)
        else:
            raise AssertionError("Production persistence must fail without preview auth or real auth.")

        settings.enable_persistence = False
        settings.enable_case_memory = True
        try:
            validate_runtime_security()
        except RuntimeError as error:
            assert "ENABLE_CASE_MEMORY=1" in str(error)
        else:
            raise AssertionError("Production case memory must fail without preview auth or real auth.")

        settings.enable_real_user_auth = True
        validate_runtime_security()
    finally:
        settings.runtime_environment = original_runtime
        settings.enable_persistence = original_persistence
        settings.enable_case_memory = original_case_memory
        settings.remind_preview_api_token = original_preview_token
        settings.enable_real_user_auth = original_real_auth
        settings.remind_allow_local_bypass = original_local_bypass

    root = Path(__file__).resolve().parents[1]
    for migration_path in (root / "supabase" / "migrations").glob("*.sql"):
        sql = migration_path.read_text(encoding="utf-8").lower()
        assert "drop table" not in sql
        assert "drop schema" not in sql
        assert "truncate " not in sql
        assert "delete from" not in sql

    assert content_hash("  a\nb  ", model="model-a") == content_hash("a b", model="model-a")
    assert content_hash("a b", model="model-a") != content_hash("a b", model="model-b")
    settings.use_stub = False
    settings.openai_api_key = None
    try:
        get_embedding_provider()
    except EmbeddingError:
        pass
    else:
        raise AssertionError("Dense retrieval must fail closed when no embedding provider is configured.")
    settings.use_stub = True

    original_cache_ttl = settings.embedding_cache_ttl_seconds
    original_cache_max = settings.embedding_cache_max_entries
    original_cache_disabled = settings.disable_embedding_cache
    original_embedding_model = settings.embedding_model
    try:
        clear_embedding_cache()
        settings.disable_embedding_cache = False
        settings.embedding_cache_ttl_seconds = 1
        settings.embedding_cache_max_entries = 2
        settings.embedding_model = "test-embedding-model-a"
        first_embedding = embed_query("same normalized query")
        assert embedding_cache_stats()["misses"] == 1
        second_embedding = embed_query(" same   normalized   query ")
        assert first_embedding == second_embedding
        assert embedding_cache_stats()["hits"] == 1
        settings.embedding_model = "test-embedding-model-b"
        embed_query("same normalized query")
        assert embedding_cache_stats()["misses"] == 2
        time.sleep(1.05)
        embed_query("same normalized query")
        assert embedding_cache_stats()["misses"] == 3
    finally:
        clear_embedding_cache()
        settings.embedding_cache_ttl_seconds = original_cache_ttl
        settings.embedding_cache_max_entries = original_cache_max
        settings.disable_embedding_cache = original_cache_disabled
        settings.embedding_model = original_embedding_model

    client = TestClient(app)
    health = client.get("/api/health")
    assert health.status_code == 200, health.text
    assert health.json() == {"status": "ok"}
    preview_headers = {"X-Remind-Preview-Token": "test-preview-token"}
    authenticated_client = TestClient(app, headers=preview_headers)

    capabilities_response = authenticated_client.get("/api/documents/capabilities")
    assert capabilities_response.status_code == 200, capabilities_response.text
    capabilities = capabilities_response.json()
    assert capabilities["docx"]["available"] is True
    assert capabilities["hwpx"]["available"] is False
    if os.name == "nt" and not require_pdf_export:
        assert capabilities["pdf"]["available"] is False, capabilities
    if require_pdf_export:
        assert capabilities["pdf"]["available"] is True, capabilities

    _run_material_upload_smoke_tests(authenticated_client)

    with tempfile.TemporaryDirectory() as temp_dir:
        os.environ["TEMP_DRAFT_DIR"] = temp_dir
        os.environ["RECOMPOSE_CACHE_DIR"] = str(Path(temp_dir) / "recompose")

        payload = {
            "case_id": "CASE-SYNTHETIC-001",
            "session_number": 5,
            "session_date": "2026-05-24",
            "counselor_name": "테스트 상담사",
            "counselor_memo": "내담자는 최근 진로 선택을 앞두고 불안과 회피가 증가했다고 보고했다.",
            "transcript_text": (
                "내담자: 지원서를 제출하려고 하면 실패할 것 같아 미루게 돼요.\n"
                "상담자: 실패할 것이라는 생각과 그때의 감정을 함께 살펴보겠습니다."
            ),
            "previous_session_summary": "이전 회기에는 불안을 유발하는 상황과 자동사고를 확인했다.",
            "target_document_type": "session_note",
            "persist": False,
        }
        missing_token = client.post("/api/notes/generate", json=payload)
        assert missing_token.status_code == 401
        invalid_token = client.post("/api/notes/generate", json=payload, headers={"X-Remind-Preview-Token": "wrong"})
        assert invalid_token.status_code == 401
        assert client.get("/api/notes/drafts").status_code == 401

        pii_text = "이름: 홍길동, 연락처 010-1234-5678, email test@example.com, 학번 2026123456"
        masked_text, pii_candidates = deidentify_text(pii_text, source="counselor_memo")
        assert "010-1234-5678" not in masked_text
        assert "test@example.com" not in masked_text
        assert "홍길동" not in masked_text
        assert "[STUDENT_ID]" in masked_text
        assert pii_candidates

        response = client.post("/api/notes/generate", json=payload, headers=preview_headers)
        assert response.status_code == 200, response.text
        data = response.json()
        assert data["session_summary_draft"]["session_info"]["case_id"] == payload["case_id"]
        assert data["session_summary_draft"]["session_info"]["session_number"] == payload["session_number"]
        assert data["session_summary_draft"]["session_content"]["text"]
        assert data["evidence_mapped_data"]["items"]
        assert data["verification_report"]["requires_counselor_review"]
        assert data["document_transform_preview"]["missing_required_fields"]
        assert data["confirmed_session_note"]["status"] == "draft_requires_counselor_confirmation"
        assert data["retrieval_report"]["enabled"] is False
        assert data["retrieved_case_context"] == []
        assert data["persistence_report"]["requested"] is False

        session_row = _build_session_row(
            SessionInput(**payload),
            GenerateNoteResponse(**data),
            user_id="test-user-id",
        )
        assert session_row["raw_input_text"] is None
        assert session_row["sanitized_input_text"]
        settings.save_raw_input = True
        raw_session_row = _build_session_row(
            SessionInput(**payload),
            GenerateNoteResponse(**data),
            user_id="test-user-id",
        )
        assert raw_session_row["raw_input_text"]
        settings.save_raw_input = False

        pii_payload = {
            **payload,
            "counselor_memo": pii_text,
            "transcript_text": "Cl: 서울시 강남구 테헤란로에 살고 010-9999-0000으로 연락 가능해요.",
        }
        pii_response = client.post("/api/notes/generate", json=pii_payload, headers=preview_headers)
        assert pii_response.status_code == 200, pii_response.text
        pii_data = pii_response.json()
        serialized_sanitized = json.dumps(pii_data["sanitized_input"], ensure_ascii=False)
        assert "010-1234-5678" not in serialized_sanitized
        assert "010-9999-0000" not in serialized_sanitized
        assert "test@example.com" not in serialized_sanitized
        assert "홍길동" not in serialized_sanitized

        memory_request = ConfirmGeneratedNoteRequest(
            note_id="00000000-0000-0000-0000-000000000011",
            confirmed_note={
                "sections": {
                    "session_theme": pii_text,
                    "client_response": "Client reported anxiety decreasing after small next actions.",
                    "next_plan": "Review student id 2026123456 only after masking.",
                }
            },
            counselor_edited=True,
            create_case_memory=True,
        )
        memory_context = ConfirmedNoteContext(
            note_id=memory_request.note_id,
            case_id=payload["case_id"],
            session_id="00000000-0000-0000-0000-000000000012",
            session_number=payload["session_number"],
            session_date=payload["session_date"],
            counselor_id="test-preview-actor",
        )
        memory_rows = _case_memory_rows_from_confirmed_note(memory_request, memory_context)
        serialized_memory_rows = json.dumps(memory_rows, ensure_ascii=False)
        assert "010-1234-5678" not in serialized_memory_rows
        assert "test@example.com" not in serialized_memory_rows
        assert "2026123456" not in serialized_memory_rows
        assert "[PHONE]" in serialized_memory_rows
        assert "[EMAIL]" in serialized_memory_rows
        assert "[STUDENT_ID]" in serialized_memory_rows
        assert all(row["source_note_id"] == memory_request.note_id for row in memory_rows)
        assert all(row["source_ref"].startswith(f"confirmed_note:{memory_request.note_id}:") for row in memory_rows)

        original_dense = settings.enable_dense_retrieval
        original_stub = settings.use_stub
        original_api_key = settings.openai_api_key
        try:
            settings.enable_dense_retrieval = True
            settings.use_stub = False
            settings.openai_api_key = None
            unavailable_rows = [row.copy() for row in memory_rows]
            assert _attach_embeddings(unavailable_rows) == 0
            assert all("embedding" not in row for row in unavailable_rows)

            settings.use_stub = True
            embedded_rows = [row.copy() for row in memory_rows]
            assert _attach_embeddings(embedded_rows) == len(embedded_rows)
            assert all(len(row["embedding"]) == settings.embedding_dimension for row in embedded_rows)
        finally:
            settings.enable_dense_retrieval = original_dense
            settings.use_stub = original_stub
            settings.openai_api_key = original_api_key

        persist_without_supabase = client.post("/api/notes/generate", json={**payload, "persist": True}, headers=preview_headers)
        assert persist_without_supabase.status_code == 200, persist_without_supabase.text
        persist_data = persist_without_supabase.json()
        assert persist_data["persistence_report"]["requested"] is True
        assert persist_data["persistence_report"]["stored"] is False

        original_enable_rag = settings.enable_rag
        original_case_retrieval = graph_nodes.retrieve_case_context
        original_case_memory_chunks = graph_nodes.retrieve_case_memory_chunks
        original_authoritative_kb_chunks = graph_nodes.retrieve_authoritative_kb_chunks
        original_template_retrieval = graph_nodes.retrieve_document_template
        original_privacy_retrieval = graph_nodes.retrieve_privacy_rules
        original_enable_dense = settings.enable_dense_retrieval
        try:
            settings.enable_rag = True

            def fake_case_context(case_id: str, current_session_id: str | None = None, max_sessions: int = 3):
                return [
                    RetrievedCaseContextItem(
                        source_ref="stored_session_note:prior-session-1",
                        session_id="prior-session-1",
                        session_number=1,
                        session_date="2026-05-01",
                        summary="이전 회기에서는 사회적 상황 불안과 회피 행동을 다룸.",
                        evidence_items=[
                            RetrievedEvidenceItem(
                                id="evidence-1",
                                source_type="direct",
                                source_ref="stored_evidence:evidence-1",
                                source_text="내담자는 지원 전 회피 행동을 보고함.",
                                linked_field="session_content",
                            )
                        ],
                    )
                ]

            def fake_template_context(target_document_type):
                return RetrievedTemplateContext(
                    target_document_type=target_document_type,
                    required_fields=["주호소", "상담 내용", "다음 계획"],
                    counselor_review_fields=["사례개념화"],
                    missing_field_checklist=["사례개념화", "목표 달성 정도"],
                    source_refs=["kb_template:session-note-demo"],
                )

            def fake_privacy_rules():
                return [
                    RetrievedPrivacyRule(
                        source_ref="kb_privacy:demo-rule",
                        title="Demo privacy rule",
                        category="privacy_rule",
                        rule="Store minimum necessary counseling data.",
                        warning="저장 전 비식별화와 동의 필요 여부를 확인하세요.",
                    )
                ]

            graph_nodes.retrieve_case_context = fake_case_context
            graph_nodes.retrieve_document_template = fake_template_context
            graph_nodes.retrieve_privacy_rules = fake_privacy_rules
            rag_response = client.post("/api/notes/generate", json=payload, headers=preview_headers)
            assert rag_response.status_code == 200, rag_response.text
            rag_data = rag_response.json()
            assert rag_data["retrieval_report"]["enabled"] is True
            assert rag_data["retrieval_report"]["case_context_count"] == 1
            assert rag_data["retrieval_report"]["template_context_found"] is True
            assert rag_data["retrieval_report"]["privacy_rule_count"] == 1
            assert rag_data["retrieved_case_context"][0]["source_ref"] == "stored_session_note:prior-session-1"
            assert rag_data["retrieved_template_context"]["missing_field_checklist"]
            assert rag_data["retrieved_privacy_context"][0]["warning"]

            def fake_case_memory_chunks(**kwargs):
                return [
                    RetrievalChunk(
                        chunk_id="case-memory-chunk-1",
                        session_id="prior-session-dense-1",
                        source_ref="case_memory:prior-session-dense-1:session_theme",
                        field_type="session_theme",
                        chunk_text="Dense prior-session memory matched career anxiety.",
                        retrieval_method="case_memory_dense",
                        similarity_score=0.82,
                        session_number=1,
                        session_date="2026-05-01",
                    )
                ]

            def fake_authoritative_kb_chunks(**kwargs):
                return [
                    RetrievalChunk(
                        chunk_id="kb-template-chunk-1",
                        document_id="kb-template-doc-1",
                        source_ref="kb:session-note-template-v1:1",
                        title="Session note template",
                        doc_category="session_note_template",
                        document_type="session_note",
                        allowed_use="documentation_structure_only",
                        authority_level="internal_demo",
                        chunk_text="Session notes require session content and next plan.",
                        retrieval_method="hybrid:dense+keyword",
                        similarity_score=0.7,
                        metadata={
                            "required_fields": ["session_content", "next_plan"],
                            "missing_field_checklist": ["next_plan"],
                        },
                    ),
                    RetrievalChunk(
                        chunk_id="kb-privacy-chunk-1",
                        document_id="kb-privacy-doc-1",
                        source_ref="kb:privacy-law-sensitive-info-demo:1",
                        title="Privacy warning",
                        doc_category="privacy_law",
                        chunk_text="Sensitive information requires consent and safety review.",
                        retrieval_method="hybrid:dense+keyword",
                        similarity_score=0.66,
                        metadata={"warning": "Review sensitive information before storage."},
                    ),
                ]

            settings.enable_dense_retrieval = True
            graph_nodes.retrieve_case_memory_chunks = fake_case_memory_chunks
            graph_nodes.retrieve_authoritative_kb_chunks = fake_authoritative_kb_chunks
            dense_response = client.post("/api/notes/generate", json=payload, headers=preview_headers)
            assert dense_response.status_code == 200, dense_response.text
            dense_data = dense_response.json()
            assert dense_data["retrieved_case_context"][0]["source_ref"].startswith("case_memory:")
            assert "kb:session-note-template-v1:1" in dense_data["retrieved_template_context"]["source_refs"]
            assert dense_data["retrieved_privacy_context"][0]["source_ref"]
        finally:
            settings.enable_rag = original_enable_rag
            settings.enable_dense_retrieval = original_enable_dense
            graph_nodes.retrieve_case_context = original_case_retrieval
            graph_nodes.retrieve_case_memory_chunks = original_case_memory_chunks
            graph_nodes.retrieve_authoritative_kb_chunks = original_authoritative_kb_chunks
            graph_nodes.retrieve_document_template = original_template_retrieval
            graph_nodes.retrieve_privacy_rules = original_privacy_retrieval

        confirm_without_persistence = client.post(
            "/api/notes/confirm",
            json={
                "note_id": "00000000-0000-0000-0000-000000000001",
                "confirmed_note": data["confirmed_session_note"],
                "counselor_edited": True,
                "create_case_memory": True,
            },
            headers=preview_headers,
        )
        assert confirm_without_persistence.status_code == 409
        spoofed_confirm = client.post(
            "/api/notes/confirm",
            json={
                "note_id": "00000000-0000-0000-0000-000000000001",
                "case_id": payload["case_id"],
                "confirmed_by": "spoofed-client",
                "confirmed_note": data["confirmed_session_note"],
            },
            headers=preview_headers,
        )
        assert spoofed_confirm.status_code == 422
        try:
            confirm_generated_note(
                ConfirmGeneratedNoteRequest(
                    note_id="00000000-0000-0000-0000-000000000001",
                    confirmed_note=data["confirmed_session_note"],
                ),
                actor="test-preview-actor",
            )
        except NoteConfirmationError as error:
            assert error.status_code == 409
        else:
            raise AssertionError("Confirmation must be server-validated and reject disabled persistence.")

        original_storage = supabase_storage_module.storage
        original_retrieval_storage = retrieval_module.storage
        original_enable_persistence = settings.enable_persistence
        original_enable_case_memory = settings.enable_case_memory
        original_supabase_url = settings.supabase_url
        original_supabase_key = settings.supabase_service_role_key
        original_enable_dense = settings.enable_dense_retrieval
        original_confirmation_enable_rag = settings.enable_rag
        try:
            fake_storage = FakeConfirmationStorage()
            supabase_storage_module.storage = fake_storage  # type: ignore[assignment]
            retrieval_module.storage = fake_storage  # type: ignore[assignment]
            settings.enable_persistence = True
            settings.enable_case_memory = True
            settings.supabase_url = "https://example.supabase.co"
            settings.supabase_service_role_key = "fake-service-key"
            settings.enable_dense_retrieval = True
            settings.enable_rag = True

            confirm_payload = ConfirmGeneratedNoteRequest(
                note_id=fake_storage.note_id,
                confirmed_note={
                    "sections": {
                        "session_theme": "Career anxiety and self-critical thoughts.",
                        "client_response": "Anxiety decreased after smaller actions.",
                        "next_plan": "Review one thought record next session.",
                    }
                },
            )
            first_confirm = confirm_generated_note(confirm_payload, actor="test-preview-actor")
            assert first_confirm.confirmation_status == "confirmed"
            assert first_confirm.memory_chunk_count == 3
            assert len(fake_storage.case_memory_chunks) == 3
            first_theme_hash = fake_storage.memory_by_field("session_theme")["content_hash"]

            fake_storage.case_memory_chunks.append(
                {
                    "id": "foreign-memory",
                    "counselor_id": "test-preview-actor",
                    "case_id": "CASE-OTHER-SYNTHETIC",
                    "session_id": "foreign-session",
                    "source_note_id": "foreign-note",
                    "field_type": "session_theme",
                    "chunk_text": "Synthetic foreign case that must never be retrieved.",
                    "source_ref": "confirmed_note:foreign-note:session_theme",
                    "embedding": embed_query("foreign case"),
                }
            )
            follow_up_context = retrieve_case_memory_chunks(
                query_text="career anxiety follow-up",
                counselor_id="test-preview-actor",
                case_id=fake_storage.case_id,
            )
            assert len(follow_up_context) == 3, [chunk.source_ref for chunk in follow_up_context]
            assert all(chunk.source_ref.startswith(f"confirmed_note:{fake_storage.note_id}:") for chunk in follow_up_context)
            assert all(chunk.chunk_text != "Synthetic foreign case that must never be retrieved." for chunk in follow_up_context)
            fake_storage.case_memory_chunks.pop()

            second_confirm = confirm_generated_note(confirm_payload, actor="test-preview-actor")
            assert second_confirm.memory_chunk_count == 3
            assert len(fake_storage.case_memory_chunks) == 3
            assert fake_storage.memory_by_field("session_theme")["content_hash"] == first_theme_hash

            try:
                confirm_generated_note(
                    ConfirmGeneratedNoteRequest(
                        note_id=fake_storage.note_id,
                        confirmed_note={
                            "sections": {
                                "session_theme": "Unmarked conflicting confirmation.",
                                "client_response": "Anxiety decreased after smaller actions.",
                                "next_plan": "Review one thought record next session.",
                            }
                        },
                        counselor_edited=False,
                    ),
                    actor="test-preview-actor",
                )
            except NoteConfirmationError as error:
                assert error.status_code == 409
            else:
                raise AssertionError("Changed reconfirmation without counselor_edited=true must be rejected.")

            fake_storage.memory_by_field("session_theme")["unchanged_marker"] = "keep"
            revised_confirm = confirm_generated_note(
                ConfirmGeneratedNoteRequest(
                    note_id=fake_storage.note_id,
                    confirmed_note={
                        "sections": {
                            "session_theme": "Revised career anxiety theme.",
                            "client_response": "Anxiety decreased after smaller actions.",
                            "next_plan": "Review one thought record next session.",
                        }
                    },
                ),
                actor="test-preview-actor",
            )
            assert revised_confirm.memory_chunk_count == 3
            assert len(fake_storage.case_memory_chunks) == 3
            assert fake_storage.memory_by_field("session_theme")["content_hash"] != first_theme_hash
            assert fake_storage.memory_by_field("session_theme")["unchanged_marker"] == "keep"
            assert fake_storage.duplicate_source_ref_groups() == 0

            try:
                confirm_generated_note(
                    ConfirmGeneratedNoteRequest(
                        note_id="00000000-0000-0000-0000-00000000ffff",
                        confirmed_note=confirm_payload.confirmed_note,
                    ),
                    actor="test-preview-actor",
                )
            except NoteConfirmationError as error:
                assert error.status_code == 404
            else:
                raise AssertionError("Nonexistent notes must be rejected.")

            fake_storage.drop_session = True
            try:
                confirm_generated_note(confirm_payload, actor="test-preview-actor")
            except NoteConfirmationError as error:
                assert error.status_code == 409
            else:
                raise AssertionError("Notes with missing sessions must be rejected.")
        finally:
            supabase_storage_module.storage = original_storage  # type: ignore[assignment]
            retrieval_module.storage = original_retrieval_storage  # type: ignore[assignment]
            settings.enable_persistence = original_enable_persistence
            settings.enable_case_memory = original_enable_case_memory
            settings.supabase_url = original_supabase_url
            settings.supabase_service_role_key = original_supabase_key
            settings.enable_dense_retrieval = original_enable_dense
            settings.enable_rag = original_confirmation_enable_rag

        save_payload = {
            "case_id": payload["case_id"],
            "session_number": payload["session_number"],
            "session_date": payload["session_date"],
            "counselor_name": payload["counselor_name"],
            "screen": "summary_draft",
            "form": payload,
            "session_topic": data["session_summary_draft"]["session_theme"]["text"],
            "visible_section_ids": ["main_issue", "session_theme", "session_content"],
            "draft_sections": [
                {
                    "id": "session_content",
                    "title": "상담 내용",
                    "content": data["session_summary_draft"]["session_content"]["text"],
                }
            ],
            "result": data,
        }
        save_response = client.post("/api/notes/drafts", json=save_payload, headers=preview_headers)
        assert save_response.status_code == 200, save_response.text
        saved = save_response.json()
        assert saved["draft_id"]
        assert saved["case_id"] == payload["case_id"]

        load_response = client.get(f"/api/notes/drafts/{saved['draft_id']}", headers=preview_headers)
        assert load_response.status_code == 200, load_response.text
        loaded = load_response.json()
        assert loaded["draft_id"] == saved["draft_id"]
        assert loaded["screen"] == "summary_draft"
        assert loaded["draft_sections"][0]["title"] == "상담 내용"

        recompose_payload = {
            "session_input": payload,
            "session_topic": "사회적 상황 불안과 평가에 대한 추측 점검",
            "visible_section_ids": ["main_issue", "session_theme", "session_content"],
        }
        first_recompose = client.post("/api/notes/recompose", json=recompose_payload, headers=preview_headers)
        assert first_recompose.status_code == 200, first_recompose.text
        first_data = first_recompose.json()
        assert first_data["cache_hit"] is False
        assert first_data["visible_section_ids"] == recompose_payload["visible_section_ids"]
        assert first_data["result"]["session_summary_draft"]["session_content"]["text"]

        second_recompose = client.post("/api/notes/recompose", json=recompose_payload, headers=preview_headers)
        assert second_recompose.status_code == 200, second_recompose.text
        second_data = second_recompose.json()
        assert second_data["cache_hit"] is True
        assert second_data["cache_key"] == first_data["cache_key"]

        supervision_payload = {
            "session_input": payload,
            "session_summary_draft": data["session_summary_draft"],
            "demo_mode": False,
            "report_date": payload["session_date"],
            "client_alias": payload["case_id"],
        }
        supervision_response = client.post("/api/notes/supervision-report", json=supervision_payload, headers=preview_headers)
        assert supervision_response.status_code == 200, supervision_response.text
        supervision = supervision_response.json()
        assert supervision["title"] == "개인상담(공개상담) 사례 수퍼비전 보고서"
        assert supervision["reportType"] == "personal_counseling_supervision"
        assert supervision["sections"]
        assert [section["id"] for section in supervision["sections"]] == [
            "A", "A-1", "A-2", "A-3", "A-4", "A-5", "A-6", "A-7", "A-8",
            "B", "B-1", "B-2", "B-3", "C", "C-1", "C-2",
        ]
        assert supervision["meta"]["institution"] == "[상담사 확인 필요]"
        assert supervision["meta"]["supervisor"] == "[상담사 확인 필요]"
        assert supervision["aiReview"]["completionChecklist"]
        assert supervision["aiReview"]["missingFields"]
        assert supervision["aiReview"]["demoInputs"] == []
        assert not any(block.get("demoValue") for section in supervision["sections"] for block in section.get("contentBlocks", []))
        assert supervision["aiReview"]["suggestedSupervisionQuestions"]

        session_export_payload = {
            "format": "docx",
            "document_type": "session_note",
            "case_id": "CASE/DEMO:*001",
            "session_number": 5,
            "session_date": "2026-05-24",
            "title": "상담 회기 기록",
            "metadata": {
                "client_alias": "CLIENT-SYNTHETIC-001",
                "counselor_name": "박상담사",
                "missing_items": ["상담 목표 표현 구체화 필요"],
                "warnings": ["근거 부족 검토 문구"],
            },
            "sections": [
                {
                    "id": "main_issue",
                    "title": "주요 호소",
                    "content": "사회적 상황에서 평가에 대한 불안과 회피를 보고함.",
                },
                {
                    "id": "session_content",
                    "title": "상담 내용",
                    "content": "첫 줄 상담 내용\n둘째 줄 상담 내용\n- 목록 항목\n최종 수정 내용이 반영됨.",
                },
            ],
        }
        docx_response = authenticated_client.post("/api/documents/export", json=session_export_payload)
        assert docx_response.status_code == 200, docx_response.text
        assert docx_response.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        filename = _download_filename(docx_response.headers["content-disposition"])
        assert filename.endswith(".docx")
        assert "CASE_DEMO__001" in filename
        assert not any(char in filename for char in '<>:"/\\|?*')
        docx_text = _extract_docx_text(docx_response.content)
        assert "상담 회기 기록" in docx_text
        assert "CLIENT-SYNTHETIC-001" in docx_text
        assert "첫 줄 상담 내용" in docx_text
        assert "둘째 줄 상담 내용" in docx_text
        assert "최종 수정 내용이 반영됨." in docx_text
        assert "missing_items" not in docx_text
        assert "상담 목표 표현 구체화 필요" not in docx_text
        assert "근거 부족 검토 문구" not in docx_text

        real_case_without_alias_payload = {
            **session_export_payload,
            "case_id": "CASE-REAL-002",
            "metadata": {"counselor_name": "박상담사"},
        }
        real_case_response = authenticated_client.post("/api/documents/export", json=real_case_without_alias_payload)
        assert real_case_response.status_code == 200, real_case_response.text
        real_case_text = _extract_docx_text(real_case_response.content)
        assert "CLIENT-SYNTHETIC-001" not in real_case_text
        assert "내담자 가명" not in real_case_text

        supervision_export_payload = {
            "format": "docx",
            "document_type": "supervision_report",
            "case_id": payload["case_id"],
            "session_number": payload["session_number"],
            "session_date": payload["session_date"],
            "title": "개인상담 사례 수퍼비전 보고서",
            "metadata": {
                "client_alias": "CLIENT-SYNTHETIC-001",
                "counselor_name": "박상담사",
                "supervisor": "이수현 상담심리사 1급",
            },
            "sections": [
                {"id": "part-c", "title": "C. 상담 과정", "level": 1},
                {
                    "id": "process",
                    "title": "C-1. 상담진행 과정 및 회기주제",
                    "level": 2,
                    "contentBlocks": [
                        {
                            "id": "paragraph-1",
                            "type": "paragraph",
                            "text": "수정된 pending edit: 불안 자동사고를 사건-생각-감정-행동으로 구분함.",
                        },
                        {
                            "id": "table-1",
                            "type": "table",
                            "rows": [
                                {"영역": "정서", "내용": "불안 80"},
                                {"영역": "행동", "내용": "지원 전 회피"},
                            ],
                        },
                        {
                            "id": "transcript-1",
                            "type": "transcript",
                            "speakerTurns": [
                                {"turnId": "t1", "speaker": "client", "text": "계속 망했다는 생각이 들어요."},
                                {"turnId": "t2", "speaker": "counselor", "text": "그 생각의 근거를 함께 보겠습니다."},
                            ],
                        },
                        {
                            "id": "reflection-1",
                            "type": "reflection_box",
                            "text": "상담자는 정서 확인과 행동 계획의 균형을 점검할 필요가 있음.",
                        },
                    ],
                },
            ],
        }
        supervision_docx_response = authenticated_client.post("/api/documents/export", json=supervision_export_payload)
        assert supervision_docx_response.status_code == 200, supervision_docx_response.text
        supervision_docx_text = _extract_docx_text(supervision_docx_response.content)
        assert "개인상담 사례 수퍼비전 보고서" in supervision_docx_text
        assert "수정된 pending edit" in supervision_docx_text
        assert "내담자: 계속 망했다는 생각이 들어요." in supervision_docx_text
        assert "영역" in supervision_docx_text
        assert "불안 80" in supervision_docx_text

        termination_export_payload = {
            "format": "docx",
            "document_type": "termination_report",
            "case_id": payload["case_id"],
            "session_number": payload["session_number"],
            "session_date": payload["session_date"],
            "title": "종결 보고서",
            "metadata": {"counselor_name": "박상담사"},
            "sections": [
                {"id": "termination_goal_process", "title": "상담 목표 및 진행 과정", "content": "진행 과정"},
                {"id": "termination_changes", "title": "주요 변화", "content": "주요 변화"},
                {"id": "termination_reason", "title": "종결 사유", "content": "합의 종결"},
                {"id": "termination_recommendation", "title": "향후 권고", "content": "향후 권고"},
                {"id": "termination_counselor_opinion", "title": "상담자 종합소견", "content": "종합소견"},
            ],
        }
        termination_docx_response = authenticated_client.post("/api/documents/export", json=termination_export_payload)
        assert termination_docx_response.status_code == 200, termination_docx_response.text
        termination_docx_text = _extract_docx_text(termination_docx_response.content)
        for expected_section in [
            "상담 목표 및 진행 과정",
            "주요 변화",
            "종결 사유",
            "향후 권고",
            "상담자 종합소견",
        ]:
            assert expected_section in termination_docx_text

        if capabilities["pdf"]["available"] or require_pdf_export:
            pdf_response = authenticated_client.post(
                "/api/documents/export",
                json={**session_export_payload, "format": "pdf"},
            )
            assert pdf_response.status_code == 200, pdf_response.text
            _assert_pdf_response(
                pdf_response.content,
                pdf_response.headers["content-type"],
                ["상담 회기 기록", "첫 줄 상담 내용"],
            )
        else:
            print(f"PDF export not exercised: {capabilities['pdf'].get('reason')}")

        invalid_format_response = authenticated_client.post(
            "/api/documents/export",
            json={**session_export_payload, "format": "xlsx"},
        )
        assert invalid_format_response.status_code == 422

        empty_sections_response = authenticated_client.post(
            "/api/documents/export",
            json={**session_export_payload, "format": "docx", "sections": []},
        )
        assert empty_sections_response.status_code == 422

        hwpx_response = authenticated_client.post("/api/documents/export", json={**session_export_payload, "format": "hwpx"})
        assert hwpx_response.status_code == 422
        assert "HWPX" in hwpx_response.text

    print(
        "Smoke test passed: health, note generation, temporary draft storage, cached recomposition, "
        "supervision report generation, document export, and material upload checks are working."
    )


class FakeConfirmationStorage:
    def __init__(self) -> None:
        self.retrieval_enabled = True
        self.note_id = "00000000-0000-0000-0000-000000000101"
        self.session_id = "00000000-0000-0000-0000-000000000102"
        self.case_id = "CASE-SYNTHETIC-001"
        self.drop_session = False
        self.generated_notes = [
            {
                "id": self.note_id,
                "case_id": self.case_id,
                "session_id": self.session_id,
                "note_type": "session_note",
                "draft_json": {"synthetic": True},
                "confirmed_json": {},
                "confirmation_status": "draft",
            }
        ]
        self.sessions = [
            {
                "id": self.session_id,
                "case_id": self.case_id,
                "session_number": 5,
                "session_date": "2026-05-24",
                "session_title": "Synthetic confirmation test",
            }
        ]
        self.cases = [{"id": self.case_id, "case_alias": self.case_id, "counselor_id": "test-preview-actor"}]
        self.case_memory_chunks: list[dict[str, object]] = []

    def maybe_single(self, table: str, query: dict[str, str | int]) -> dict[str, object] | None:
        rows = self.select(table, query)
        return rows[0] if rows else None

    def select(self, table: str, query: dict[str, str | int]) -> list[dict[str, object]]:
        if table == "generated_notes":
            return self._filter_by_eq(self.generated_notes, "id", str(query.get("id") or ""))
        if table == "sessions":
            if self.drop_session:
                return []
            return self._filter_by_eq(self.sessions, "id", str(query.get("id") or ""))
        if table == "cases":
            return self._filter_by_eq(self.cases, "id", str(query.get("id") or ""))
        if table == "case_memory_chunks":
            source_note_id = self._eq_value(str(query.get("source_note_id") or ""))
            return [row for row in self.case_memory_chunks if row.get("source_note_id") == source_note_id]
        return []

    def update(
        self,
        table: str,
        values: dict[str, object],
        *,
        query: dict[str, str | int],
        return_representation: bool = True,
    ) -> list[dict[str, object]]:
        rows = self.select(table, query)
        for row in rows:
            row.update(values)
        return rows if return_representation else []

    def upsert(self, table: str, rows: list[dict[str, object]], *, on_conflict: str) -> list[dict[str, object]]:
        assert table == "case_memory_chunks"
        assert on_conflict == "source_note_id,field_type"
        result = []
        for row in rows:
            existing = next(
                (
                    current
                    for current in self.case_memory_chunks
                    if current.get("source_note_id") == row.get("source_note_id")
                    and current.get("field_type") == row.get("field_type")
                ),
                None,
            )
            if existing:
                existing.update(row)
                result.append(existing)
            else:
                stored = dict(row)
                self.case_memory_chunks.append(stored)
                result.append(stored)
        return result

    def rpc(self, function_name: str, params: dict[str, object]) -> list[dict[str, object]]:
        if function_name == "match_case_memory_chunks":
            rows = [
                row
                for row in self.case_memory_chunks
                if row.get("counselor_id") == params.get("filter_counselor_id")
                and row.get("case_id") == params.get("filter_case_id")
            ]
            return [
                {
                    **row,
                    "chunk_id": row.get("id") or f"chunk-{index}",
                    "similarity_score": 0.9 - index * 0.01,
                    "retrieval_method": "case_memory_dense",
                    "metadata": row.get("metadata_json") or {},
                }
                for index, row in enumerate(rows[: int(params.get("match_count") or 5)])
            ]
        if function_name == "log_retrieval_event":
            return []
        raise AssertionError(f"Unexpected RPC: {function_name}")

    def insert(self, table: str, rows: list[dict[str, object]]) -> list[dict[str, object]]:
        assert table == "retrieval_logs"
        return rows

    def memory_by_field(self, field_type: str) -> dict[str, object]:
        for row in self.case_memory_chunks:
            if row.get("field_type") == field_type:
                return row
        raise AssertionError(f"Missing memory row for {field_type}")

    def duplicate_source_ref_groups(self) -> int:
        counts: dict[object, int] = {}
        for row in self.case_memory_chunks:
            counts[row.get("source_ref")] = counts.get(row.get("source_ref"), 0) + 1
        return sum(1 for count in counts.values() if count > 1)

    def _filter_by_eq(self, rows: list[dict[str, object]], key: str, condition: str) -> list[dict[str, object]]:
        value = self._eq_value(condition)
        return [row for row in rows if row.get(key) == value]

    @staticmethod
    def _eq_value(condition: str) -> str:
        return condition[3:] if condition.startswith("eq.") else condition


if __name__ == "__main__":
    main()
