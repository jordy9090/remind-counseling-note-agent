"""Document export service for DOCX/PDF output."""
from __future__ import annotations

import html
import os
import re
import unicodedata
from abc import ABC, abstractmethod
from dataclasses import dataclass
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from string import Template
from typing import Iterable
from urllib.parse import quote

from app.schemas.document import DocumentContentBlock, DocumentExportRequest, DocumentSection


DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_CONTENT_TYPE = "application/pdf"

DOCUMENT_TYPE_LABELS = {
    "session_note": "상담일지",
    "supervision_report": "수퍼비전보고서",
    "termination_report": "종결보고서",
}

REVIEW_NOTICE = (
    "윤리·개인정보 보호 기준을 준수하여 작성합니다."
)

SUPERVISION_PLACEHOLDER = "[상담사 확인 필요]"

BLOCKED_METADATA_KEYS = {
    "missing_items",
    "warnings",
    "unsupported_claims",
    "unsupported_or_risky_claims",
    "ai_review",
    "aiReview",
    "needs_human_review",
    "completion_checklist",
}


class DocumentExportError(Exception):
    """Base class for export errors that should become client-facing responses."""


class UnsupportedExportFormat(DocumentExportError):
    """Raised when the requested export format is intentionally unavailable."""


class DocumentExportValidationError(DocumentExportError):
    """Raised when an export request has no renderable document body."""


class DocumentExportRuntimeError(DocumentExportError):
    """Raised when the server cannot render an otherwise valid export."""


@dataclass(frozen=True)
class DocumentExportResult:
    content: bytes
    content_type: str
    filename: str

    @property
    def headers(self) -> dict[str, str]:
        ascii_fallback = build_ascii_filename_fallback(self.filename)
        encoded = quote(self.filename)
        return {
            "Content-Disposition": (
                f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{encoded}"
            )
        }


class DocumentExporter(ABC):
    """Base exporter interface; HWPX can plug in here later."""

    content_type: str
    extension: str

    @abstractmethod
    def export(self, request: DocumentExportRequest) -> bytes:
        """Return the rendered file as bytes."""


class DocxDocumentExporter(DocumentExporter):
    content_type = DOCX_CONTENT_TYPE
    extension = "docx"

    def export(self, request: DocumentExportRequest) -> bytes:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(14)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)

        font_name = os.getenv("REMIND_DOCX_FONT_FAMILY", "Malgun Gothic")
        configure_docx_styles(document, font_name)
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_paragraph.add_run(REVIEW_NOTICE)
        footer_run.font.size = Pt(7.5)
        footer_run.font.color.rgb = RGBColor(71, 85, 105)
        set_run_font(footer_run, font_name)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(4)
        title.paragraph_format.space_after = Pt(10)
        add_paragraph_border(title, "111827", 14)
        title_run = title.add_run(request.title)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(15, 23, 42)
        set_run_font(title_run, font_name)

        supervision_grid = request.document_type == "supervision_report"
        info_rows = build_supervision_metadata_grid(request) if supervision_grid else build_metadata_rows(request)
        table = document.add_table(rows=len(info_rows), cols=4 if supervision_grid else 2)
        table.style = "Table Grid"
        if supervision_grid:
            set_table_column_widths(table, [25, 55, 35, 55])
        for index, info_row in enumerate(info_rows):
            cells = table.rows[index].cells
            for column_index, value in enumerate(info_row):
                cells[column_index].text = value
                if column_index % 2 == 0:
                    shade_cell(cells[column_index], "E5E7EB")
                    for paragraph in cells[column_index].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            set_run_font(run, font_name)
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        set_run_font(run, font_name)

        document.add_paragraph()

        for doc_section in renderable_sections(request.sections):
            if doc_section.level <= 1:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.keep_with_next = True
                run = paragraph.add_run(doc_section.title)
                run.bold = True
                run.font.size = Pt(14)
                set_run_font(run, font_name)
            else:
                heading = document.add_heading(doc_section.title, level=min(max(doc_section.level, 2), 3))
                heading.paragraph_format.keep_with_next = True
                for run in heading.runs:
                    set_run_font(run, font_name)

            if doc_section.content_blocks:
                for block in doc_section.content_blocks:
                    add_block_to_docx(document, block, font_name)
            else:
                add_content_to_docx(document, doc_section.content, font_name)

        # Word requires a paragraph after a trailing table. Keep that structural
        # paragraph tiny so a final reflection box cannot create a blank page.
        end_marker = document.add_paragraph()
        end_marker.paragraph_format.space_before = Pt(0)
        end_marker.paragraph_format.space_after = Pt(0)
        end_marker.paragraph_format.line_spacing = Pt(1)
        marker_run = end_marker.add_run(" ")
        marker_run.font.size = Pt(1)
        marker_run.font.hidden = True
        set_run_font(marker_run, font_name)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()


class PdfDocumentExporter(DocumentExporter):
    content_type = PDF_CONTENT_TYPE
    extension = "pdf"

    def export(self, request: DocumentExportRequest) -> bytes:
        try:
            from weasyprint import HTML
            from weasyprint.text.fonts import FontConfiguration

            html_text = render_pdf_html(request)
            font_config = FontConfiguration()
            return HTML(string=html_text, url_fetcher=block_external_resource).write_pdf(font_config=font_config)
        except (ImportError, OSError) as error:
            raise DocumentExportRuntimeError(
                "PDF 렌더링 런타임을 불러오지 못했습니다. WeasyPrint와 Pango/GObject 시스템 라이브러리 "
                "설치를 확인해주세요."
            ) from error


class HwpxDocumentExporter(DocumentExporter):
    content_type = "application/vnd.hancom.hwpx"
    extension = "hwpx"

    def export(self, request: DocumentExportRequest) -> bytes:
        raise UnsupportedExportFormat(
            "HWPX 내보내기는 검증된 HWPX 템플릿 ZIP 구조가 준비된 뒤 활성화할 수 있습니다. "
            "현재는 Word/한글 호환 DOCX 또는 PDF를 선택해주세요."
        )


class DocumentExportService:
    def __init__(self) -> None:
        self._exporters: dict[str, DocumentExporter] = {
            "docx": DocxDocumentExporter(),
            "pdf": PdfDocumentExporter(),
            "hwpx": HwpxDocumentExporter(),
        }

    def export(self, request: DocumentExportRequest) -> DocumentExportResult:
        exporter = self._exporters[request.format]
        sections = list(renderable_sections(request.sections))
        if not sections:
            raise DocumentExportValidationError("내보낼 수 있는 본문 섹션이 없습니다.")

        content = exporter.export(request.model_copy(update={"sections": sections}))
        return DocumentExportResult(
            content=content,
            content_type=exporter.content_type,
            filename=build_download_filename(request, exporter.extension),
        )

    def capabilities(self) -> dict[str, dict[str, str | bool | None]]:
        pdf_available, pdf_reason = check_pdf_runtime()
        return {
            "docx": {"available": True, "reason": None},
            "pdf": {"available": pdf_available, "reason": None if pdf_available else pdf_reason},
            "hwpx": {
                "available": False,
                "reason": "Verified HWPX template is not configured.",
            },
        }


def build_metadata_rows(request: DocumentExportRequest) -> list[tuple[str, str]]:
    rows = [
        ("문서 유형", DOCUMENT_TYPE_LABELS.get(request.document_type, request.document_type)),
        ("사례 ID", request.case_id),
        ("회기", f"{request.session_number}회기"),
        ("날짜", request.session_date or "미기재"),
    ]
    for key, value in request.metadata.items():
        if key in BLOCKED_METADATA_KEYS:
            continue
        if value is None or value == "":
            continue
        rows.append((humanize_metadata_key(key), stringify_cell_value(value)))
    return rows


def build_supervision_metadata_grid(request: DocumentExportRequest) -> list[tuple[str, str, str, str]]:
    metadata = request.metadata
    return [
        (
            "상담자",
            stringify_cell_value(metadata.get("counselor_name") or ""),
            "소속 상담기관",
            stringify_cell_value(metadata.get("institution") or ""),
        ),
        (
            "수퍼바이저",
            stringify_cell_value(metadata.get("supervisor") or ""),
            "수퍼비전 일시 및 장소",
            stringify_cell_value(metadata.get("supervision_date_place") or ""),
        ),
    ]


@lru_cache(maxsize=1)
def check_pdf_runtime() -> tuple[bool, str | None]:
    try:
        from weasyprint import HTML
        from weasyprint.text.fonts import FontConfiguration

        HTML(string="<html><body><p>pdf capability check</p></body></html>").write_pdf(
            font_config=FontConfiguration(),
        )
    except Exception:
        return False, "WeasyPrint native runtime is unavailable."
    return True, None


def humanize_metadata_key(key: str) -> str:
    labels = {
        "client_alias": "내담자 가명",
        "counselor_name": "상담자",
        "institution": "소속 상담기관",
        "supervisor": "수퍼바이저",
        "supervision_date_place": "수퍼비전 일시 및 장소",
        "report_date": "보고서 기준일",
    }
    return labels.get(key, key.replace("_", " "))


def stringify_cell_value(value: object) -> str:
    if isinstance(value, list):
        return ", ".join(stringify_cell_value(item) for item in value)
    if isinstance(value, dict):
        return ", ".join(f"{key}: {stringify_cell_value(item)}" for key, item in value.items())
    text = str(value)
    return "" if text.strip() == SUPERVISION_PLACEHOLDER else text


def renderable_sections(sections: Iterable[DocumentSection]) -> Iterable[DocumentSection]:
    for section in sections:
        if section_has_content(section):
            yield section


def section_has_content(section: DocumentSection) -> bool:
    if section.level <= 1 and section.title.strip():
        return True
    if isinstance(section.content, str) and section.content.strip():
        return True
    if isinstance(section.content, list) and any(str(item).strip() for item in section.content):
        return True
    return any(block_has_content(block) for block in section.content_blocks)


def block_has_content(block: DocumentContentBlock) -> bool:
    if block.text and block.text.strip():
        return True
    if block.rows:
        return True
    return any(turn.text.strip() for turn in block.speaker_turns)


def add_content_to_docx(document, content: str | list[str] | None, font_name: str) -> None:
    if isinstance(content, list):
        for item in content:
            add_docx_text_line(document, str(item), font_name, force_bullet=True)
        return

    for line in split_preserving_blank_lines(clean_export_text(content or "")):
        add_docx_text_line(document, line, font_name)


def add_docx_text_line(document, line: str, font_name: str, force_bullet: bool = False) -> None:
    text = line.strip()
    if not text:
        document.add_paragraph()
        return
    bullet_match = re.match(r"^\s*[-*•]\s+(.+)$", line)
    number_match = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
    if force_bullet or bullet_match:
        paragraph = document.add_paragraph(style="List Bullet")
        text = bullet_match.group(1) if bullet_match else text
    elif number_match:
        paragraph = document.add_paragraph(style="List Number")
        text = number_match.group(1)
    else:
        paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, font_name)


def add_block_to_docx(document, block: DocumentContentBlock, font_name: str) -> None:
    if block.label:
        label_paragraph = document.add_paragraph()
        label_paragraph.paragraph_format.keep_with_next = True
        label_run = label_paragraph.add_run(block.label)
        label_run.bold = True
        set_run_font(label_run, font_name)
    if block.type == "table" and block.rows:
        add_table_block_to_docx(document, block.rows, font_name)
        return
    if block.type == "transcript" and block.speaker_turns:
        for index, turn in enumerate(block.speaker_turns, 1):
            label = "내담자" if turn.speaker == "client" else "상담자" if turn.speaker == "counselor" else "화자"
            silence = f" (침묵 {turn.silence_seconds}초)" if turn.silence_seconds is not None else ""
            add_docx_text_line(document, f"{index}. {label}: {turn.text}{silence}", font_name)
        return
    if block.type == "reflection_box":
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        shade_cell(table.cell(0, 0), "F3F4F6")
        table.cell(0, 0).text = clean_export_text(block.text or "")
        for paragraph in table.cell(0, 0).paragraphs:
            for run in paragraph.runs:
                set_run_font(run, font_name)
        return
    add_content_to_docx(document, block.text or "", font_name)


def add_table_block_to_docx(document, rows: list[dict[str, object]], font_name: str) -> None:
    headers = ordered_table_headers(rows)
    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    if len(headers) == 6:
        set_table_column_widths(table, [15, 22, 20, 43, 38, 32])
    mark_table_header_repeat(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = str(header)
        shade_cell(cell, "E2E8F0")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                set_run_font(run, font_name)
    for row_index, row in enumerate(rows, start=1):
        for column_index, header in enumerate(headers):
            cell = table.cell(row_index, column_index)
            cell.text = stringify_cell_value(row.get(header, ""))
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name)
    document.add_paragraph()


def ordered_table_headers(rows: list[dict[str, object]]) -> list[str]:
    headers: list[str] = []
    for row in rows:
        for key in row.keys():
            if key not in headers:
                headers.append(str(key))
    return headers or ["내용"]


def configure_docx_styles(document, font_name: str) -> None:
    from docx.shared import Pt, RGBColor

    for style_name in ("Normal", "Body Text"):
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = font_name
            style.font.size = Pt(10.5)
            set_style_font(style, font_name)
    for style_name in ("Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        if style_name in document.styles:
            style = document.styles[style_name]
            set_style_font(style, font_name)
            if style_name.startswith("Heading"):
                style.font.color.rgb = RGBColor(15, 23, 42)
                style.font.bold = True
                style.paragraph_format.space_before = Pt(10 if style_name == "Heading 2" else 7)
                style.paragraph_format.space_after = Pt(5)


def set_table_column_widths(table, widths_mm: list[int]) -> None:
    from docx.shared import Mm

    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths_mm):
            if index < len(row.cells):
                row.cells[index].width = Mm(width)


def mark_table_header_repeat(row) -> None:
    from docx.oxml import OxmlElement

    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))


def add_paragraph_border(paragraph, color: str, size: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), color)
        borders.append(border)
    p_pr.append(borders)


def set_style_font(style, font_name: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, font_name: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def render_pdf_html(request: DocumentExportRequest) -> str:
    template_path = Path(__file__).resolve().parents[1] / "templates" / "document_export.html"
    template = Template(template_path.read_text(encoding="utf-8"))
    font_family = sanitize_css_font_family(os.getenv("REMIND_PDF_FONT_FAMILY", "Malgun Gothic"))
    return template.safe_substitute(
        title=escape_text(request.title),
        font_family=font_family,
        metadata_rows=render_metadata_rows_html(request),
        sections=render_sections_html(request.sections),
        review_notice=escape_text(REVIEW_NOTICE),
    )


def render_metadata_rows_html(request: DocumentExportRequest) -> str:
    if request.document_type == "supervision_report":
        return "\n".join(
            "<tr>" + "".join(
                f"<{('th' if index % 2 == 0 else 'td')}>{escape_text(value)}</{('th' if index % 2 == 0 else 'td')}>"
                for index, value in enumerate(row)
            ) + "</tr>"
            for row in build_supervision_metadata_grid(request)
        )
    rows = []
    for label, value in build_metadata_rows(request):
        rows.append(
            "<tr>"
            f"<th>{escape_text(label)}</th>"
            f"<td>{escape_text(value)}</td>"
            "</tr>"
        )
    return "\n".join(rows)


def render_sections_html(sections: Iterable[DocumentSection]) -> str:
    rendered = []
    for section in sections:
        heading_tag = "h2" if section.level <= 2 else "h3"
        body = (
            "".join(render_block_html(block) for block in section.content_blocks)
            if section.content_blocks
            else render_content_html(section.content)
        )
        rendered.append(
            "<section class=\"document-section\">"
            f"<{heading_tag}>{escape_text(section.title)}</{heading_tag}>"
            f"{body}"
            "</section>"
        )
    return "\n".join(rendered)


def render_content_html(content: str | list[str] | None) -> str:
    if isinstance(content, list):
        items = "".join(f"<li>{escape_text(str(item))}</li>" for item in content if str(item).strip())
        return f"<ul>{items}</ul>" if items else ""

    parts: list[str] = []
    for line in split_preserving_blank_lines(clean_export_text(content or "")):
        if not line.strip():
            parts.append("<div class=\"blank-line\"></div>")
            continue
        bullet_match = re.match(r"^\s*[-*•]\s+(.+)$", line)
        number_match = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if bullet_match:
            parts.append(f"<ul><li>{escape_text(bullet_match.group(1))}</li></ul>")
        elif number_match:
            parts.append(f"<ol><li>{escape_text(number_match.group(1))}</li></ol>")
        else:
            parts.append(f"<p>{escape_text(line)}</p>")
    return "\n".join(parts)


def render_block_html(block: DocumentContentBlock) -> str:
    label_html = f'<p class="block-label">{escape_text(block.label)}</p>' if block.label else ""
    if block.type == "table" and block.rows:
        headers = ordered_table_headers(block.rows)
        header_html = "".join(f"<th>{escape_text(header)}</th>" for header in headers)
        rows_html = []
        for row in block.rows:
            cells = "".join(f"<td>{escape_text(stringify_cell_value(row.get(header, '')))}</td>" for header in headers)
            rows_html.append(f"<tr>{cells}</tr>")
        return f"{label_html}<table class=\"content-table\"><thead><tr>{header_html}</tr></thead><tbody>{''.join(rows_html)}</tbody></table>"
    if block.type == "transcript" and block.speaker_turns:
        turns = []
        for index, turn in enumerate(block.speaker_turns, 1):
            label = "내담자" if turn.speaker == "client" else "상담자" if turn.speaker == "counselor" else "화자"
            silence = f" (침묵 {turn.silence_seconds}초)" if turn.silence_seconds is not None else ""
            turns.append(
                "<div class=\"transcript-turn\">"
                f"<strong>{index}. {escape_text(label)}</strong>"
                f"<span>{escape_text(turn.text + silence)}</span>"
                "</div>"
            )
        return f"{label_html}<div class=\"transcript\">{''.join(turns)}</div>"
    if block.type == "reflection_box":
        return f"{label_html}<div class=\"reflection-box\">{render_content_html(block.text or '')}</div>"
    return label_html + render_content_html(block.text or "")


def escape_text(value: object) -> str:
    return html.escape(str(value), quote=True)


def clean_export_text(value: object) -> str:
    text = str(value)
    return "" if text.strip() == SUPERVISION_PLACEHOLDER else text


def sanitize_css_font_family(value: str) -> str:
    families = []
    for family in value.split(","):
        clean = re.sub(r"[^0-9A-Za-z가-힣 _.-]", "", family).strip()
        if clean:
            families.append(f'"{clean}"')
    families.extend(['"Apple SD Gothic Neo"', '"Noto Sans CJK KR"', '"Noto Sans KR"', "sans-serif"])
    return ", ".join(families)


def block_external_resource(url: str, *args, **kwargs):
    raise ValueError(f"External resources are disabled for document export: {url}")


def split_preserving_blank_lines(value: str) -> list[str]:
    return value.replace("\r\n", "\n").replace("\r", "\n").split("\n")


def build_download_filename(request: DocumentExportRequest, extension: str) -> str:
    document_label = DOCUMENT_TYPE_LABELS.get(request.document_type, request.document_type)
    parts = [
        sanitize_filename_part(document_label, "document"),
        sanitize_filename_part(request.case_id, "case"),
        sanitize_filename_part(f"{request.session_number}회기", "session"),
        sanitize_filename_part(request.session_date or "date", "date"),
    ]
    return "_".join(parts) + f".{extension}"


def sanitize_filename_part(value: object, fallback: str) -> str:
    text = unicodedata.normalize("NFC", str(value)).strip()
    text = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", text)
    text = re.sub(r"\s+", "_", text)
    text = text.strip("._ ")
    return text[:80] or fallback


def build_ascii_filename_fallback(filename: str) -> str:
    fallback = unicodedata.normalize("NFKD", filename).encode("ascii", "ignore").decode("ascii")
    fallback = re.sub(r"[^A-Za-z0-9_.-]+", "_", fallback).strip("._")
    if not fallback or "." not in fallback:
        extension = filename.rsplit(".", 1)[-1] if "." in filename else "bin"
        fallback = f"document_export.{extension}"
    return fallback
