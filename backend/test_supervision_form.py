"""Focused regression tests for the Korean Counseling Psychological Association form."""
from __future__ import annotations

import unittest
from io import BytesIO
from pathlib import Path

from docx import Document

from app.graph.supervision_report import PLACEHOLDER, run_supervision_report_pipeline
from app.schemas.document import DocumentExportRequest
from app.schemas.note import SessionInput, SupervisionReportRequest
from app.services.document_export import DocxDocumentExporter, PdfDocumentExporter, render_pdf_html


EXPECTED_SECTION_IDS = [
    "A", "A-1", "A-2", "A-3", "A-4", "A-5", "A-6", "A-7", "A-8",
    "B", "B-1", "B-2", "B-3", "C", "C-1", "C-2",
]


def request_for(**updates) -> SupervisionReportRequest:
    payload = {
        "case_id": "CASE-SYNTHETIC-001",
        "session_number": 5,
        "session_date": "2026-08-22",
        "counselor_name": "테스트 상담사",
        "counselor_memo": "진로 선택을 앞두고 불안과 회피가 증가했다고 보고함.",
        "transcript_text": "내담자: 실패할 것 같아 미루게 돼요.\n상담자: 그 생각을 함께 살펴보겠습니다.",
        "previous_session_summary": "불안 상황과 자동사고를 확인함.",
        "target_document_type": "session_note",
        "persist": False,
    }
    values = {"session_input": SessionInput(**payload), **updates}
    return SupervisionReportRequest(**values)


class SupervisionFormTests(unittest.TestCase):
    def test_official_title_and_section_order(self):
        report = run_supervision_report_pipeline(request_for())
        self.assertEqual(report.title, "개인상담(공개상담) 사례 수퍼비전 보고서")
        self.assertEqual([section.id for section in report.sections], EXPECTED_SECTION_IDS)

    def test_missing_information_is_not_invented(self):
        report = run_supervision_report_pipeline(request_for())
        sections = {section.id: section for section in report.sections}
        for section_id in ("A-1", "A-2", "A-4", "A-5", "A-8", "B-1", "B-3"):
            block = sections[section_id].contentBlocks[0]
            self.assertEqual(block.text, PLACEHOLDER)
            self.assertEqual(block.evidenceStatus, "missing")
            self.assertTrue(block.missingInputs)

    def test_goals_and_strategy_are_separate_editable_blocks(self):
        report = run_supervision_report_pipeline(request_for(
            agreed_counseling_goal="합의 목표",
            clinical_counseling_goal="임상 목표",
            counseling_strategy="구체적 개입 전략",
        ))
        section = next(item for item in report.sections if item.id == "B-2")
        self.assertEqual(
            [(block.label, block.text) for block in section.contentBlocks],
            [
                ("내담자와 합의한 목표", "합의 목표"),
                ("상담자의 임상적 목표", "임상 목표"),
                ("상담전략", "구체적 개입 전략"),
            ],
        )

    def test_progress_table_and_full_transcript_preserve_structure(self):
        request = request_for(
            maximum_sessions=12,
            transcript_mode="full",
            session_events=[{
                "session_number": 2,
                "session_date": "2026-05-01",
                "duration_minutes": 40,
                "topic": "불안 상황 탐색",
                "attendance_status": "late",
                "attendance_reason": "교통 지연",
            }],
            previous_supervisions=[{"supervision_date": "2026-04-20", "feedback": "언어반응을 점검함"}],
        )
        request.session_input.transcript_text = "내담자: 말하기 어렵습니다. (침묵 5초)\n상담자: 천천히 말씀해도 됩니다."
        report = run_supervision_report_pipeline(request)
        sections = {section.id: section for section in report.sections}
        progress_blocks = sections["C-1"].contentBlocks
        self.assertEqual(progress_blocks[0].text, "12회기")
        self.assertEqual(list(progress_blocks[1].rows[0]), ["회기", "상담일", "소요시간", "회기 주제", "취소·지각·결석", "이전 수퍼비전"])
        transcript = sections["C-2"].contentBlocks[0]
        self.assertEqual(transcript.type, "transcript")
        self.assertEqual(len(transcript.speakerTurns), 2)
        self.assertEqual(transcript.speakerTurns[0].silenceSeconds, 5)

    def test_ai_organized_clinical_language_stays_tentative(self):
        section = lambda text: {"text": text, "evidence_type": "inferred", "source_refs": []}
        report = run_supervision_report_pipeline(request_for(session_summary_draft={
            "session_info": {
                "case_id": "CASE-TEST", "client_alias": "내담자", "session_number": 5,
                "session_date": "2026-08-22", "counselor_name": "상담자",
            },
            "session_theme": section("회피 패턴을 확인하였다."),
            "presenting_problem": section("불안과 관련되어 있다고 판단하였다."),
            "session_content": section("급성 위험도는 낮은 수준으로 판단하였다."),
            "counselor_intervention": section("개입 효과가 있다고 판단하였다."),
            "client_response": section("반응이 안정적이라고 판단하였다."),
            "reflection": section("종결이 가능하다고 판단하였다."),
            "next_plan": section("다음 회기 계획"),
        }))
        rendered = " ".join(
            [block.text or "" for item in report.sections for block in item.contentBlocks]
            + [value for item in report.sections for block in item.contentBlocks for row in (block.rows or []) for value in row.values()]
        )
        self.assertNotIn("판단하였다", rendered)
        self.assertIn("추후 확인이 필요하다", rendered)

    def test_docx_and_pdf_use_official_metadata_grid(self):
        export_request = DocumentExportRequest(
            format="docx",
            document_type="supervision_report",
            case_id="CASE-TEST",
            session_number=1,
            session_date="2026-08-22",
            title="개인상담(공개상담) 사례 수퍼비전 보고서",
            metadata={"counselor_name": "상담자", "institution": "기관"},
            sections=[{"id": "A", "title": "A. 내담자 기본 정보", "level": 1}],
        )
        document = Document(BytesIO(DocxDocumentExporter().export(export_request)))
        self.assertEqual(len(document.tables[0].rows), 2)
        self.assertEqual(len(document.tables[0].columns), 4)
        self.assertEqual(document.tables[0].cell(0, 0).text, "상담자")
        self.assertEqual(document.tables[0].cell(1, 2).text, "수퍼비전 일시 및 장소")
        self.assertEqual(document.tables[0].cell(1, 3).text, "")
        self.assertIn("개인정보 보호 기준", document.sections[0].footer.paragraphs[0].text)
        self.assertFalse(any("개인정보 보호 기준" in paragraph.text for paragraph in document.paragraphs))
        html = render_pdf_html(export_request)
        self.assertIn("개인상담(공개상담) 사례 수퍼비전 보고서", html)
        self.assertIn("<th>소속 상담기관</th>", html)

        pdf = PdfDocumentExporter().export(export_request.model_copy(update={"format": "pdf"}))
        self.assertTrue(pdf.startswith(b"%PDF-"))
        self.assertGreater(len(pdf), 1_000)


if __name__ == "__main__":
    unittest.main()
